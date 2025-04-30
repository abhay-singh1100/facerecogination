from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os
from datetime import datetime

def create_project_report():
    doc = Document()
    
    # Title Page
    title = doc.add_heading('Face Recognition Attendance System', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add project details
    doc.add_paragraph().add_run('\nProject Report').bold = True
    doc.add_paragraph(f'Date: {datetime.now().strftime("%B %d, %Y")}')
    doc.add_paragraph('Developed by: [Your Name]')
    doc.add_paragraph('Guide: [Guide Name]')
    
    # Table of Contents
    doc.add_page_break()
    doc.add_heading('Table of Contents', level=1)
    sections = [
        "1. Introduction",
        "2. System Requirements",
        "3. Technology Stack",
        "4. System Architecture",
        "5. Features and Functionality",
        "6. Implementation Details",
        "7. Security Measures",
        "8. Testing and Validation",
        "9. Future Enhancements",
        "10. Conclusion"
    ]
    for section in sections:
        doc.add_paragraph(section, style='List Number')
    
    # Introduction
    doc.add_page_break()
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph('The Face Recognition Attendance System is an advanced biometric solution designed to automate the process of attendance tracking using facial recognition technology. This system provides a contactless, efficient, and secure method of recording attendance, making it particularly relevant in modern educational institutions and organizations.')
    
    # System Requirements
    doc.add_heading('2. System Requirements', level=1)
    doc.add_paragraph('Hardware Requirements:')
    hardware_reqs = [
        "Computer/Server with minimum 4GB RAM",
        "Webcam with minimum 720p resolution",
        "Storage space for face database",
        "Internet connectivity for web interface"
    ]
    for req in hardware_reqs:
        doc.add_paragraph(req, style='List Bullet')
    
    doc.add_paragraph('Software Requirements:')
    software_reqs = [
        "Python 3.11 or higher",
        "OpenCV for image processing",
        "dlib for face recognition",
        "Flask web framework",
        "Modern web browser"
    ]
    for req in software_reqs:
        doc.add_paragraph(req, style='List Bullet')
    
    # Technology Stack
    doc.add_heading('3. Technology Stack', level=1)
    tech_stack = {
        'Backend': ['Python', 'Flask', 'OpenCV', 'dlib', 'face_recognition'],
        'Frontend': ['HTML5', 'CSS3', 'JavaScript'],
        'Database': ['CSV File System'],
        'Authentication': ['Session-based authentication']
    }
    for category, technologies in tech_stack.items():
        doc.add_paragraph(f'{category}:', style='Heading 2')
        for tech in technologies:
            doc.add_paragraph(tech, style='List Bullet')
    
    # System Architecture
    doc.add_heading('4. System Architecture', level=1)
    doc.add_paragraph('The system follows a client-server architecture with the following components:')
    components = [
        "Web Interface (Client)",
        "Flask Server (Backend)",
        "Face Recognition Engine",
        "Database Management",
        "Authentication System"
    ]
    for component in components:
        doc.add_paragraph(component, style='List Bullet')
    
    # Features and Functionality
    doc.add_heading('5. Features and Functionality', level=1)
    features = [
        "Real-time face detection and recognition",
        "Liveness detection to prevent spoofing",
        "User registration with face enrollment",
        "Automatic attendance marking",
        "Administrative dashboard",
        "Attendance reports and analytics",
        "Export functionality for attendance data",
        "Secure authentication system"
    ]
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')
    
    # Implementation Details
    doc.add_heading('6. Implementation Details', level=1)
    doc.add_paragraph('The system implementation involves several key components:')
    
    implementation_details = {
        'Face Detection': 'Uses MediaPipe for initial face detection with high accuracy',
        'Face Recognition': 'Employs dlib\'s face recognition model based on deep learning',
        'Liveness Detection': 'Implements blink detection using MediaPipe Face Mesh',
        'Web Interface': 'Responsive design using modern web technologies',
        'Database': 'Efficient CSV-based storage system for attendance records'
    }
    for topic, detail in implementation_details.items():
        p = doc.add_paragraph()
        p.add_run(f'{topic}: ').bold = True
        p.add_run(detail)
    
    # Security Measures
    doc.add_heading('7. Security Measures', level=1)
    security_features = [
        "Anti-spoofing measures through liveness detection",
        "Secure storage of face encodings",
        "Session-based authentication",
        "Input validation and sanitization",
        "CSRF protection",
        "Secure file handling"
    ]
    for feature in security_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    # Testing and Validation
    doc.add_heading('8. Testing and Validation', level=1)
    testing_approaches = {
        'Unit Testing': 'Individual component testing',
        'Integration Testing': 'Testing component interactions',
        'System Testing': 'End-to-end system validation',
        'Performance Testing': 'System performance under various conditions'
    }
    for approach, description in testing_approaches.items():
        p = doc.add_paragraph()
        p.add_run(f'{approach}: ').bold = True
        p.add_run(description)
    
    # Future Enhancements
    doc.add_heading('9. Future Enhancements', level=1)
    enhancements = [
        "Integration with existing ERP systems",
        "Mobile application development",
        "Enhanced analytics and reporting",
        "Multi-factor authentication",
        "Cloud deployment options"
    ]
    for enhancement in enhancements:
        doc.add_paragraph(enhancement, style='List Bullet')
    
    # Conclusion
    doc.add_heading('10. Conclusion', level=1)
    doc.add_paragraph('The Face Recognition Attendance System successfully demonstrates the application of modern computer vision and web technologies in creating an efficient and secure attendance management solution. The system provides a robust foundation for future enhancements and can be adapted for various organizational needs.')
    
    # Save the document
    doc.save('Face_Recognition_Attendance_System_Report.docx')
    print("Project report has been generated successfully!")

if __name__ == "__main__":
    create_project_report()